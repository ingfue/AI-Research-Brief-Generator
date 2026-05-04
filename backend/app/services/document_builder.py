"""
Word template filler for the Research Brief.

Loads the template .docx, fills Table 0 metadata cells, replaces the
instructional paragraphs under each heading with AI-generated content
(including markdown-to-Word formatting for bold, italic, sub-headings,
bullet lists, and tables), and returns the filled document as bytes.
"""

import io
import json
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from app.models.schemas import SectionName, MetadataFields

TEMPLATE_PATH = Path(__file__).resolve().parent.parent.parent / "templates" / "research_brief_template.docx"

HEADING_SECTION_MAP: dict[str, list[SectionName]] = {
    "1. Client & Brand": [SectionName.CLIENT_BRAND],
    "2. Project Overview / Background": [SectionName.PROJECT_OVERVIEW],
    "3. Objectives": [SectionName.OBJECTIVES],
    "4. Research Questions": [SectionName.RESEARCH_QUESTIONS],
    "Data Analysis Timeframe": [SectionName.DATA_TIMEFRAME],
    "How this Research Will Be Used": [SectionName.RESEARCH_USAGE],
    "Deliverables": [SectionName.DELIVERABLES],
    "6. Project Timeline": [SectionName.TIMELINE],
    "7. Key Assumptions": [SectionName.KEY_ASSUMPTIONS],
    "8. Additional Information": [SectionName.ADDITIONAL_INFO],
}

_INLINE_RE = re.compile(
    r"(\*\*\*(.+?)\*\*\*"   # ***bold italic***
    r"|\*\*(.+?)\*\*"        # **bold**
    r"|\*(.+?)\*)"           # *italic*
)


def _is_heading(paragraph) -> bool:
    style_name = paragraph.style.name if paragraph.style else ""
    return style_name.startswith("Heading")


def _clear_paragraph(paragraph):
    """Remove all runs from a paragraph, keeping the paragraph element itself."""
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)
    p_elem = paragraph._element
    for child in list(p_elem):
        if child.tag == qn("w:r"):
            p_elem.remove(child)


def _add_inline_formatting(paragraph, text: str):
    """Parse **bold**, *italic*, ***bold-italic*** and create properly
    formatted runs on *paragraph*.  Plain text segments become normal runs."""
    last_end = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > last_end:
            paragraph.add_run(text[last_end:m.start()])

        if m.group(2):
            run = paragraph.add_run(m.group(2))
            run.bold = True
            run.italic = True
        elif m.group(3):
            run = paragraph.add_run(m.group(3))
            run.bold = True
        elif m.group(4):
            run = paragraph.add_run(m.group(4))
            run.italic = True

        last_end = m.end()

    if last_end < len(text):
        paragraph.add_run(text[last_end:])


def _safe_style(doc: Document, preferred: str, fallback: str = "Normal"):
    """Return *preferred* style name if it exists in the document, else *fallback*.

    Iterates the style collection by name to avoid the deprecated style_id
    lookup path.
    """
    available = {s.name for s in doc.styles}
    return preferred if preferred in available else fallback


def _is_table_row(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|") and stripped.count("|") >= 3


def _apply_table_borders(tbl_element):
    """Add thin single-line borders to every cell so the table is visible."""
    NSMAP = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    tbl_pr = tbl_element.find(qn("w:tblPr"))
    if tbl_pr is None:
        from lxml import etree
        tbl_pr = etree.SubElement(tbl_element, qn("w:tblPr"))

    from lxml import etree
    borders = etree.SubElement(tbl_pr, qn("w:tblBorders"))
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        etree.SubElement(borders, qn(f"w:{edge}"), attrib={
            qn("w:val"): "single",
            qn("w:sz"): "4",
            qn("w:space"): "0",
            qn("w:color"): "000000",
        })


def _create_table_element(doc: Document, table_lines: list[str]):
    """Parse markdown table lines into a Word table and return its XML element."""
    rows_data: list[list[str]] = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if all(re.match(r"^[-:]+$", c) for c in cells if c):
            continue
        rows_data.append(cells)

    if not rows_data:
        return None

    num_cols = max(len(r) for r in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=num_cols)

    _apply_table_borders(table._element)

    for row_idx, cells in enumerate(rows_data):
        for col_idx in range(num_cols):
            cell_text = cells[col_idx].strip() if col_idx < len(cells) else ""
            cell = table.rows[row_idx].cells[col_idx]
            _clear_paragraph(cell.paragraphs[0])
            _add_inline_formatting(cell.paragraphs[0], cell_text)
            if row_idx == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True

    tbl_element = table._element
    tbl_element.getparent().remove(tbl_element)
    return tbl_element


def _build_section_paragraphs(doc: Document, content: str, insert_after_element):
    """Insert markdown-aware paragraphs (and tables) after *insert_after_element*.

    Handles: ### sub-headings, **bold** / *italic* inline formatting,
    ``- bullet`` lists, ``| col | col |`` markdown tables, and plain
    paragraphs.
    """
    lines = content.split("\n")
    parent = insert_after_element.getparent()
    insert_index = list(parent).index(insert_after_element) + 1
    offset = 0

    i = 0
    while i < len(lines):
        line = lines[i]

        if not line.strip():
            i += 1
            continue

        # --- horizontal rule — skip ---
        if line.strip() in ("---", "***", "___"):
            i += 1
            continue

        # --- markdown table block ---
        if _is_table_row(line):
            table_lines: list[str] = []
            while i < len(lines) and _is_table_row(lines[i]):
                table_lines.append(lines[i])
                i += 1
            el = _create_table_element(doc, table_lines)
            if el is not None:
                parent.insert(insert_index + offset, el)
                offset += 1
            continue

        # --- sub-heading (## / ### / ####) ---
        heading_match = re.match(r"^(#{2,4})\s+(.*)", line.strip())
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            style = _safe_style(doc, f"Heading {min(level, 4)}")
            para = doc.add_paragraph(style=style)
            _add_inline_formatting(para, text)
            el = para._element
            parent.remove(el)
            parent.insert(insert_index + offset, el)
            offset += 1
            i += 1
            continue

        # --- bullet / list item ---
        bullet_match = re.match(r"^(\s*)[-*]\s+(.*)", line)
        if bullet_match:
            indent = len(bullet_match.group(1))
            text = bullet_match.group(2)
            style_name = "List Bullet 2" if indent >= 2 else "List Bullet"
            style = _safe_style(doc, style_name)
            para = doc.add_paragraph(style=style)
            _add_inline_formatting(para, text)
            el = para._element
            parent.remove(el)
            parent.insert(insert_index + offset, el)
            offset += 1
            i += 1
            continue

        # --- regular paragraph ---
        para = doc.add_paragraph(style=_safe_style(doc, "Normal"))
        _add_inline_formatting(para, line.strip())
        el = para._element
        parent.remove(el)
        parent.insert(insert_index + offset, el)
        offset += 1
        i += 1


class DocumentBuilder:
    def build(
        self,
        metadata: MetadataFields,
        sections: dict[SectionName, str],
    ) -> bytes:
        """
        Fill the Word template with generated content and return as bytes.

        Args:
            metadata: Structured metadata for the header table.
            sections: Dict mapping SectionName -> generated text content.
        """
        doc = Document(str(TEMPLATE_PATH))

        self._fill_metadata_table(doc, metadata)
        self._fill_sections(doc, sections)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def _fill_metadata_table(self, doc: Document, metadata: MetadataFields):
        """Fill the first table (7 rows x 2 cols) with metadata values."""
        if not doc.tables:
            return

        table = doc.tables[0]
        field_map = {
            0: metadata.project_name,
            1: metadata.client,
            2: metadata.client_contact,
            3: metadata.additional_stakeholders,
            4: metadata.version,
            5: metadata.hours_allocation,
            6: metadata.prepared_by,
        }

        for row_idx, value in field_map.items():
            if row_idx < len(table.rows) and value:
                cell = table.rows[row_idx].cells[1]
                # Clear existing text and set new value
                for paragraph in cell.paragraphs:
                    _clear_paragraph(paragraph)
                cell.paragraphs[0].add_run(value)

    def _fill_sections(self, doc: Document, sections: dict[SectionName, str]):
        """
        Walk through the document paragraphs, identify headings, and replace
        the instructional paragraphs beneath each heading with agent-generated
        content.
        """
        paragraphs = doc.paragraphs
        i = 0

        while i < len(paragraphs):
            para = paragraphs[i]
            if not _is_heading(para):
                i += 1
                continue

            heading_text = para.text.strip()
            section_names = HEADING_SECTION_MAP.get(heading_text)

            if section_names is None:
                i += 1
                continue

            # Collect content from all mapped sections
            combined_content = ""
            for sn in section_names:
                content = sections.get(sn, "")
                if content:
                    combined_content += content + "\n"
            combined_content = combined_content.strip()

            if not combined_content:
                i += 1
                continue

            # Find all 'normal' paragraphs between this heading and the next heading
            j = i + 1
            paras_to_remove = []
            while j < len(paragraphs):
                if _is_heading(paragraphs[j]):
                    break
                paras_to_remove.append(j)
                j += 1

            # Remove instructional paragraphs (in reverse to preserve indices)
            for idx in reversed(paras_to_remove):
                p_elem = paragraphs[idx]._element
                p_elem.getparent().remove(p_elem)

            # Insert new content paragraphs after the heading
            _build_section_paragraphs(doc, combined_content, para._element)

            # Re-read paragraphs since we modified the document
            paragraphs = doc.paragraphs
            i += 1

    def build_from_raw(
        self,
        metadata_json: str,
        sections: dict[str, str],
    ) -> bytes:
        """
        Convenience method that accepts raw string data.

        Args:
            metadata_json: JSON string of MetadataFields.
            sections: Dict mapping section name strings to content.
        """
        metadata = MetadataFields(**json.loads(metadata_json))
        typed_sections = {}
        for key, value in sections.items():
            try:
                typed_sections[SectionName(key)] = value
            except ValueError:
                continue
        return self.build(metadata, typed_sections)
